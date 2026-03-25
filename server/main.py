import os
import shutil
import subprocess
import uuid
from html import escape
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
import fitz
from pdf2docx import Converter
from PIL import Image
from pydantic import BaseModel
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from starlette.background import BackgroundTask
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

load_dotenv()

app = FastAPI(title="Python File Converter API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).resolve().parent
TMP_DIR = BASE_DIR / "tmp"
TMP_DIR.mkdir(parents=True, exist_ok=True)

MAX_FILE_SIZE = 10 * 1024 * 1024
MAX_EDITOR_CONTENT_LENGTH = 100000

SUPPORTED_CONVERSIONS = {
    "pdf-to-word": {
        "input_extensions": {".pdf"},
        "output_extension": ".docx",
    },
    "word-to-pdf": {
        "input_extensions": {".doc", ".docx"},
        "output_extension": ".pdf",
    },
    "image-to-pdf": {
        "input_extensions": {".jpg", ".jpeg", ".png"},
        "output_extension": ".pdf",
    },
    "pdf-to-text": {
        "input_extensions": {".pdf"},
        "output_extension": ".txt",
    },
    "word-to-text": {
        "input_extensions": {".doc", ".docx"},
        "output_extension": ".txt",
    },
    "pdf-to-image": {
        "input_extensions": {".pdf"},
        "output_extension": ".png",
    },
    "image-to-word": {
        "input_extensions": {".jpg", ".jpeg", ".png"},
        "output_extension": ".docx",
    },
}


def safe_delete(path: Path) -> None:
    try:
        if path.exists():
            path.unlink()
    except OSError:
        pass


def cleanup_files(*paths: Path) -> None:
    for file_path in paths:
        safe_delete(file_path)


def get_conversion_config(conversion_type: str) -> dict:
    config = SUPPORTED_CONVERSIONS.get(conversion_type)
    if not config:
        raise HTTPException(status_code=400, detail="Invalid conversion type selected.")
    return config


async def save_upload_to_temp(upload_file: UploadFile) -> Path:
    extension = Path(upload_file.filename or "").suffix.lower()
    temp_input_path = TMP_DIR / f"{uuid.uuid4().hex}{extension}"

    total_size = 0
    with temp_input_path.open("wb") as output:
        while True:
            chunk = await upload_file.read(1024 * 1024)
            if not chunk:
                break
            total_size += len(chunk)
            if total_size > MAX_FILE_SIZE:
                safe_delete(temp_input_path)
                raise HTTPException(status_code=400, detail="File too large. Maximum size is 10MB.")
            output.write(chunk)

    await upload_file.close()
    return temp_input_path


def convert_pdf_to_word(input_path: Path, output_path: Path) -> None:
    converter = Converter(str(input_path))
    try:
        converter.convert(str(output_path))
    finally:
        converter.close()


def convert_word_to_pdf(input_path: Path, output_path: Path) -> None:
    output_dir = output_path.parent
    command = [
        "soffice",
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir),
        str(input_path),
    ]

    try:
        subprocess.run(command, check=True, capture_output=True, text=True)
    except FileNotFoundError as exc:
        raise HTTPException(
            status_code=500,
            detail="LibreOffice is required for Word to PDF conversion. Install LibreOffice and ensure 'soffice' is in PATH.",
        ) from exc
    except subprocess.CalledProcessError as exc:
        message = exc.stderr.strip() or "Word to PDF conversion failed."
        raise HTTPException(status_code=500, detail=message) from exc

    generated_pdf = output_dir / f"{input_path.stem}.pdf"
    if not generated_pdf.exists():
        raise HTTPException(status_code=500, detail="Word to PDF conversion did not produce an output file.")

    if generated_pdf != output_path:
        shutil.move(str(generated_pdf), str(output_path))


def convert_image_to_pdf(input_path: Path, output_path: Path) -> None:
    with Image.open(input_path) as image:
        rgb_image = image.convert("RGB")
        rgb_image.save(output_path, "PDF", resolution=100.0)


def convert_pdf_to_image(input_path: Path, output_path: Path) -> None:
    with fitz.open(str(input_path)) as pdf_doc:
        if len(pdf_doc) == 0:
            raise HTTPException(status_code=400, detail="PDF has no pages.")

        first_page = pdf_doc[0]
        pix = first_page.get_pixmap(matrix=fitz.Matrix(2, 2))
        pix.save(str(output_path))


def convert_image_to_word(input_path: Path, output_path: Path) -> None:
    document = Document()
    document.add_heading("Image Document", level=1)
    document.add_picture(str(input_path), width=Inches(6.2))
    document.save(str(output_path))


def write_text_file(content: str, output_path: Path) -> None:
    output_path.write_text(content or "", encoding="utf-8")


def text_to_html(content: str) -> str:
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if not lines:
        return "<p></p>"
    return "".join(f"<p>{escape(line)}</p>" for line in lines)


def html_to_plain_text(content_html: str) -> str:
    soup = BeautifulSoup(content_html, "html.parser")
    return soup.get_text("\n", strip=True)


def extract_block_nodes(soup: BeautifulSoup):
    blocks = soup.find_all(["p", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6"])
    if blocks:
        return blocks
    return [soup]


def create_docx_from_html(content_html: str, output_path: Path, title: str) -> None:
    document = Document()
    document.add_heading(title, level=1)

    soup = BeautifulSoup(content_html, "html.parser")
    for block in extract_block_nodes(soup):
        text = block.get_text(" ", strip=True)
        if not text:
            continue
        document.add_paragraph(text)

    document.save(str(output_path))


def create_pdf_from_html(content_html: str, output_path: Path, title: str) -> None:
    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    styles = getSampleStyleSheet()

    story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]

    soup = BeautifulSoup(content_html, "html.parser")
    for block in extract_block_nodes(soup):
        safe_text = escape(block.get_text(" ", strip=True)) or "&nbsp;"
        story.append(Paragraph(safe_text, styles["BodyText"]))
        story.append(Spacer(1, 8))

    doc.build(story)


def extract_text_from_pdf(input_path: Path) -> str:
    text_chunks = []
    with fitz.open(str(input_path)) as pdf_doc:
        for page in pdf_doc:
            text_chunks.append(page.get_text("text"))
    return "\n".join(text_chunks).strip()


def extract_text_from_docx(input_path: Path) -> str:
    document = Document(str(input_path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    return "\n".join(paragraphs).strip()


def convert_doc_to_docx(input_path: Path, output_dir: Path) -> Path:
    command = [
        "soffice",
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(output_dir),
        str(input_path),
    ]

    try:
        subprocess.run(command, check=True, capture_output=True, text=True)
    except FileNotFoundError as exc:
        raise HTTPException(
            status_code=500,
            detail="LibreOffice is required to import .doc files. Install LibreOffice and ensure 'soffice' is in PATH.",
        ) from exc
    except subprocess.CalledProcessError as exc:
        message = exc.stderr.strip() or "Importing .doc file failed."
        raise HTTPException(status_code=500, detail=message) from exc

    converted_path = output_dir / f"{input_path.stem}.docx"
    if not converted_path.exists():
        raise HTTPException(status_code=500, detail=".doc to .docx conversion did not produce an output file.")

    return converted_path


class EditorExportRequest(BaseModel):
    content: Optional[str] = ""
    contentHtml: Optional[str] = ""
    format: str
    title: Optional[str] = "Untitled Document"


@app.get("/api/health")
def health_check() -> JSONResponse:
    return JSONResponse({"ok": True, "message": "Python server is running"})


@app.post("/api/convert")
async def convert_file(file: UploadFile = File(...), conversionType: str = Form(...)):
    config = get_conversion_config(conversionType)
    original_extension = Path(file.filename or "").suffix.lower()

    if original_extension not in config["input_extensions"]:
        allowed = ", ".join(sorted(config["input_extensions"]))
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type for {conversionType}. Allowed: {allowed}",
        )

    input_path = await save_upload_to_temp(file)
    output_path = TMP_DIR / f"{uuid.uuid4().hex}{config['output_extension']}"
    converted_docx_path = None

    try:
        if conversionType == "pdf-to-word":
            convert_pdf_to_word(input_path, output_path)
        elif conversionType == "word-to-pdf":
            convert_word_to_pdf(input_path, output_path)
        elif conversionType == "image-to-pdf":
            convert_image_to_pdf(input_path, output_path)
        elif conversionType == "pdf-to-text":
            content = extract_text_from_pdf(input_path)
            write_text_file(content, output_path)
        elif conversionType == "word-to-text":
            if original_extension == ".docx":
                content = extract_text_from_docx(input_path)
            else:
                converted_docx_path = convert_doc_to_docx(input_path, TMP_DIR)
                content = extract_text_from_docx(converted_docx_path)
            write_text_file(content, output_path)
        elif conversionType == "pdf-to-image":
            convert_pdf_to_image(input_path, output_path)
        elif conversionType == "image-to-word":
            convert_image_to_word(input_path, output_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported conversion type.")
    except HTTPException:
        cleanup_files(input_path, output_path)
        if converted_docx_path:
            cleanup_files(converted_docx_path)
        raise
    except Exception as exc:
        cleanup_files(input_path, output_path)
        if converted_docx_path:
            cleanup_files(converted_docx_path)
        raise HTTPException(status_code=500, detail=f"Conversion failed: {exc}") from exc

    download_name = f"converted-{uuid.uuid4().hex[:8]}{config['output_extension']}"
    cleanup_targets = [input_path, output_path]
    if converted_docx_path:
        cleanup_targets.append(converted_docx_path)

    return FileResponse(
        path=output_path,
        filename=download_name,
        media_type="application/octet-stream",
        background=BackgroundTask(cleanup_files, *cleanup_targets),
    )


@app.post("/api/editor/export")
def export_editor_content(payload: EditorExportRequest):
    export_format = payload.format.lower()
    content_html = (payload.contentHtml or "").strip()
    if not content_html:
        fallback_text = (payload.content or "").strip()
        content_html = text_to_html(fallback_text)

    content_text = html_to_plain_text(content_html)
    title = (payload.title or "Untitled Document").strip() or "Untitled Document"

    if export_format not in {"pdf", "docx"}:
        raise HTTPException(status_code=400, detail="Invalid export format. Choose 'pdf' or 'docx'.")

    if not content_text:
        raise HTTPException(status_code=400, detail="Editor content cannot be empty.")

    if len(content_text) > MAX_EDITOR_CONTENT_LENGTH:
        raise HTTPException(status_code=400, detail="Editor content is too long.")

    output_extension = ".pdf" if export_format == "pdf" else ".docx"
    output_path = TMP_DIR / f"{uuid.uuid4().hex}{output_extension}"

    try:
        if export_format == "pdf":
            create_pdf_from_html(content_html, output_path, title)
        else:
            create_docx_from_html(content_html, output_path, title)
    except Exception as exc:
        cleanup_files(output_path)
        raise HTTPException(status_code=500, detail=f"Editor export failed: {exc}") from exc

    download_name = f"{title.replace(' ', '-').lower()}-{uuid.uuid4().hex[:8]}{output_extension}"
    return FileResponse(
        path=output_path,
        filename=download_name,
        media_type="application/octet-stream",
        background=BackgroundTask(cleanup_files, output_path),
    )


@app.post("/api/editor/import")
async def import_editor_content(file: UploadFile = File(...)):
    extension = Path(file.filename or "").suffix.lower()
    if extension not in {".pdf", ".doc", ".docx"}:
        raise HTTPException(status_code=400, detail="Only PDF and Word files (.doc, .docx) are supported for editing.")

    input_path = await save_upload_to_temp(file)
    converted_docx_path = None

    try:
        if extension == ".pdf":
            content = extract_text_from_pdf(input_path)
        elif extension == ".docx":
            content = extract_text_from_docx(input_path)
        else:
            converted_docx_path = convert_doc_to_docx(input_path, TMP_DIR)
            content = extract_text_from_docx(converted_docx_path)
    except HTTPException:
        cleanup_files(input_path)
        if converted_docx_path:
            cleanup_files(converted_docx_path)
        raise
    except Exception as exc:
        cleanup_files(input_path)
        if converted_docx_path:
            cleanup_files(converted_docx_path)
        raise HTTPException(status_code=500, detail=f"Failed to import file for editing: {exc}") from exc

    cleanup_files(input_path)
    if converted_docx_path:
        cleanup_files(converted_docx_path)

    title = Path(file.filename or "Untitled Document").stem
    return JSONResponse(
        {
            "title": title,
            "content": content,
            "htmlContent": text_to_html(content),
            "sourceType": extension.lstrip("."),
        }
    )


@app.exception_handler(HTTPException)
async def http_exception_handler(_, exc: HTTPException):
    return JSONResponse(status_code=exc.status_code, content={"error": exc.detail})


if __name__ == "__main__":
    import uvicorn

    port = int(os.getenv("PORT", "5000"))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)
